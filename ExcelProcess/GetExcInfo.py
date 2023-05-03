# -*- coding: UTF-8 -*-
"""
@Introduce: 提取论文中的描述符信息
    输入文件：'D:\yxs\课程\大四\毕设\Data\描述符简介表格.docx'
    输出文件：'描述符信息表.xlsx'
    注意：运行完程序后，还需手动添加 粗粒度层、细粒度层、概念层 到excel表格中

@Project ：PythonSever
@File ：GetExcInfo.py
@Author ：小小小松
@Date ：2023/3/20 11:16
"""
import os
import re

import docx
import docxlatex
import pandas as pd

from data._init_ import DataPath


def docxFunc(path):
    """
    使用docx库提取含有word中的文本信息。

    注意：docx库不能处理带公式的word文本

    :param path: 待处理word文件的路径
    :return: None
    """
    # 得到文本对象
    document = docx.Document(path)
    # 打印输出每一个单元格
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
            print('\n')


def docLatexFunc(path):
    """
    使用doclatex库处理带公式的word文本，返回处理后的信息

    :param path: 待处理word文件的路径
    :return: 列表形式text，存储了表格中的所有信息包括：中文描述符、英文描述符、简介、来源、公式
    """
    document = docxlatex.Document(path)
    text = document.get_text()
    text = text.replace('属性1：“', '').replace('属性2：“', '').replace('”', '')
    text = text.split('\n')
    while '' in text:
        text.remove('')
    # print(text[-1])
    return text


def getDocxInfo(text):
    """
     根据text列表，由正则表达式提取出每个英文描述符所在列表中的索引（位置）index，并构建结果集

    :param text: --string数组形式：存储docx库提取出的列表信息
    :return: insertList--List<dict> 每条数据存储在每条字典中，所有字典以列表形式组成。
    """

    reslut = []
    insertIndex = []
    # 正则表达式规则：匹配大小写英文字母和空格，用来筛选描述符
    pattern = re.compile(r'[-A-Za-z0-9,()\s]+')
    # 匹配出描述符名字，存储描述符名字所在索引
    for i in range(len(text)):
        str = text[i]
        res = pattern.match(str)
        if (res is not None) and (str == res.group()):
            # print(i - 1, text[i - 1])
            insertIndex.append(i - 1)

    # 把text分成一个一个小组，得到无序
    for i in range(len(insertIndex)):
        insertList = []
        localIndex = insertIndex[i]
        if i + 1 == len(insertIndex):
            nextIndex = len(text)
        else:
            nextIndex = insertIndex[i + 1]
        while localIndex < nextIndex:
            insertList.append(text[localIndex])
            localIndex = localIndex + 1

        dic = {'NodeName': insertList[1], 'ZhName': insertList[0], 'Introduce': None, 'Source': None, 'Formula': None}
        formulaList = []
        for enum in insertList[2:]:
            if r'简介：' in enum:
                dic['Introduce'] = enum.replace('简介：', '')
            elif r'来源：' in enum:
                dic['Source'] = enum.replace('来源：', '')
            elif r'公式' in enum:
                formulaList.append(enum)
        if len(formulaList) != 0:
            if len(formulaList) == 1:
                dic['Formula'] = formulaList[0]
        # print(dic)
        reslut.append(dic)
    return reslut


if __name__ == '__main__':
    # 构建word路径
    word = "描述符简介表格.docx"
    WordPath = os.path.join(DataPath, word)
    # 得到源数据列表
    text = docLatexFunc(WordPath)
    # print(text)

    insertList = getDocxInfo(text)
    # print(insertList)
    df = pd.DataFrame(insertList)
    # print(df)
    # 不需要添加序号
    df.to_excel('描述符信息表.xlsx', index=False)
